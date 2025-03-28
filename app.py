import datetime
import pandas as pd
import docx
import zipfile
import re
import os
import xml.etree.ElementTree as ET
 
from docx.oxml.ns import qn  # ✅ Import for handling Word namespaces

def load_config_from_excel(excel_path, sheet_name):
    """Load configuration for a given Word document from an Excel sheet."""
    df = pd.read_excel(excel_path, sheet_name=sheet_name, engine="openpyxl")
    config = df.set_index("Parameter")["Value"].to_dict()

    # Convert comma-separated sections into a list
    if "Sections" in config:
        config["Sections"] = [s.strip() for s in str(config["Sections"]).split(",")]

    return config

def extract_text_by_page(doc):
    """Extract text page-wise from the document."""
    text_by_page = {}
    page_number = 1
    text_by_page[page_number] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_by_page[page_number].append(text)
        # If page number is detected in text, assume a new page
        if text.startswith("Page ") and text.split()[-1].isdigit():
            page_number += 1
            text_by_page[page_number] = []

    return text_by_page

# def extract_section_names(doc):
#     """Extract section names by checking for headings and bold text inside runs."""
#     sections = []
#     for para in doc.paragraphs:
#         # Check if paragraph is styled as a Heading
#         if para.style.name.startswith("Heading"):
#             sections.append(para.text.strip())
#         # Check if any part of the paragraph is bold
#         elif any(run.bold for run in para.runs):
#             sections.append(para.text.strip())
#         # Check if paragraph is written in all uppercase (possible section heading)
#         elif para.text.isupper() and len(para.text) > 3:  # Ignore short words
#             sections.append(para.text.strip())
#     return sections

import re

def get_numbering_text(para):
    """Extracts numbering from a paragraph if it exists"""
    numbering = para._element.find(".//" + qn("w:numPr"))  # ✅ Use `qn("w:numPr")`
    if numbering is not None:
        numbering_text = []
        for run in para.runs:
            numbering_text.append(run.text)
        return " ".join(numbering_text).strip()
    return None

def extract_section_names(doc):
    """Extracts section names from the document, including numbered headings"""
    section_names = []
    
    for para in doc.paragraphs:
        text = para.text.strip()

        # ✅ Ignore empty lines or non-section content
        if not text:
            continue

        # ✅ Handle both numbered and unnumbered headings
        numbering = get_numbering_text(para)
        if numbering:
            text = numbering + " " + text  # Preserve numbered format

        # ✅ Remove unwanted TOC or page number artifacts (e.g., '\t6' at the end)
        text = re.sub(r'\t\d+$', '', text)

        section_names.append(text)

    return section_names

def extract_toc_sections(doc):
    """Extract section names from the Table of Contents (TOC)"""
    toc_sections = []

    for para in doc.paragraphs:
        text = para.text.strip()
        
        # ✅ Ignore empty paragraphs
        if not text:
            continue
        
        # ✅ TOC lines usually have page numbers at the end (e.g., "7.1 Assumptions 16")
        match = re.match(r"([\d\.]+)?\s*(.+?)\s+\d+$", text)  # Extracts "7.1 Assumptions" without the page number
        
        if match:
            section_name = match.group(2).strip()  # Extract only the section name
            toc_sections.append(section_name)

    return toc_sections

def validate_sections_using_toc(doc, config_sections):
    """Validate sections using the extracted TOC instead of document body"""
    extracted_sections = extract_toc_sections(doc)
    
    # ✅ Normalize for comparison
    def normalize(text):
        return re.sub(r'[^a-zA-Z0-9 ]', '', text).strip().lower()
    
    expected_sections = {normalize(sec) for sec in config_sections}
    extracted_sections = {normalize(sec) for sec in extracted_sections}

    missing_sections = expected_sections - extracted_sections
    unexpected_sections = extracted_sections - expected_sections

    return missing_sections, unexpected_sections


def extract_table_content(doc):
    """Extract key-value pairs from tables, handling multi-line values."""
    table_data = {}
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip().lower().replace(" ", "")
                value = row.cells[1].text.strip().replace("\n", " ")
                table_data[key] = value

                # 🔥 Debug: Print extracted keys to compare
                # print(f"Extracted Table Key: {key} | Value: {value}")

    return table_data



def extract_page1_content(doc):
    """Extract key-value pairs from both tables and paragraphs on Page 1"""
    page1_content = {}
    
    # ✅ Extract from normal paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        match = re.match(r"(.+?):\s*(.+)", text)  # Looks for "Key: Value"
        if match:
            key, value = match.groups()
            page1_content[key.strip().lower()] = value.strip()

    # ✅ Extract from tables
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:  # Ensures key-value format
                key = row.cells[0].text.strip().lower()
                value = row.cells[1].text.strip()
                page1_content[key] = value

    return page1_content

# def extract_table_content(doc):
#     """Extract key-value pairs from tables, handling multi-line values."""
#     table_data = {}
#     for table in doc.tables:
#         for row in table.rows:
#             if len(row.cells) >= 2:  # Ensure we have at least key-value format
#                 key = row.cells[0].text.strip().lower().replace(" ", "")
#                 value = row.cells[1].text.strip()
#                 table_data[key] = value
#     return table_data

def extract_table_content_fixed(doc):
    """Extract key-value pairs from tables, handling merged text in a single cell."""
    table_data = {}

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            
            if len(cells) == 1:  # 🛑 Entire row is in a single cell
                multi_line_content = cells[0].split("\n")  # Split by new lines
                for line in multi_line_content:
                    match = re.match(r"(.+?):\s*(.+)", line)  # Detect "Key: Value"
                    if match:
                        key, value = match.groups()
                        table_data[key.lower().replace(" ", "")] = value.strip()

            elif len(cells) == 2:  # ✅ Proper key-value format
                key, value = cells
                table_data[key.lower().replace(" ", "")] = value.strip()

    return table_data

def validate_page_1_content(text_by_page, doc, config):
    """Validate key-value pairs on Page 1, checking both text and tables."""
    mismatches = {}

    debug_tables_first_page(doc)

    # ✅ Extract plain text from Page 1
    page_1_text = ""
    if 1 in text_by_page:
        page_1_text = " ".join(text_by_page[1]).lower().replace(" ", "")

    # ✅ Extract key-value pairs from tables
    # table_data = extract_table_content(doc)
    # ✅ Extract key-value pairs from tables (now handling merged cells)
    table_data = extract_table_content_fixed(doc)

    # ✅ Validate expected values from config
    for key, expected_value in config.items():
        if key.startswith("Page_1_"):
            normalized_expected = str(expected_value).lower().replace(" ", "")

            # 🔥 Debug: Print expected values
            print(f"Validating {key}: '{normalized_expected}'")

            found_in_text = normalized_expected in page_1_text
            found_in_table = any(
                normalized_expected == v.lower().replace(" ", "") for v in table_data.values()
            )

            if not (found_in_text or found_in_table):
                mismatches[key] = f"Expected '{expected_value}' not found"

    return mismatches


def debug_tables(doc):
    """Prints out table structure to debug missing content."""
    if not doc.tables:
        print("❌ No tables found in the document.")
    else:
        print(f"✅ Found {len(doc.tables)} tables.")


def debug_tables_first_page(doc):
    """Prints tables that appear on the first page before the first page break."""
    found_tables = False
    for table_idx, table in enumerate(doc.tables):
        # print(f"\n🔍 Table {table_idx + 1} Content:")
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            # print(f"  ➝ Row {row_idx + 1}: {row_data}")
        found_tables = True
        # print("-" * 50)

    if not found_tables:
        print("❌ No tables found on the first page.")

# def extract_document_revision_history(doc):
#     """Extracts 'Document Revision History' section from paragraphs and tables."""
#     revision_text = ""

#     # ✅ Search in paragraphs
#     for para in doc.paragraphs:
#         if "document revision history" in para.text.lower():
#             revision_text += para.text + "\n"

#     # ✅ Search in tables (since it might be stored as a table)
#     for table in doc.tables:
#         for row in table.rows:
#             row_text = " | ".join([cell.text.strip() for cell in row.cells])
#             if "document revision history" in row_text.lower():
#                 revision_text += row_text + "\n"

#     return revision_text.strip()

# def extract_document_revision_history_from_table(doc):
#     """Extracts the 'Document Revision History' table and returns it as a dictionary."""
#     for table in doc.tables:
#         if table.rows and "revision number" in table.rows[0].cells[0].text.lower():
#             headers = [cell.text.strip().lower().replace(" ", "_") for cell in table.rows[0].cells]  # Normalize headers
#             table_data = []

#             for row in table.rows[1:]:  # Skip header row
#                 row_data = {headers[i]: cell.text.strip() for i, cell in enumerate(row.cells)}
#                 if any(row_data.values()):  # Only add non-empty rows
#                     table_data.append(row_data)

#             return table_data  # List of dictionaries with row data
    
#     return None  # Returns None if no matching table is found

def extract_document_revision_history_from_table(doc):
    """Extracts 'Document Revision History' table correctly with original column names."""
    for table in doc.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells]

            # ✅ Identify the header row by checking key terms
            if "Revision Number" in row_values and "Revision Date" in row_values:
                # headers = row_values  # ✅ Keep the original case
                headers = [h.lower() for h in row_values]  # ✅ Convert to lowercase with underscores
                
                data = []

                for r in table.rows[1:]:  # ✅ Process data rows
                    values = [cell.text.strip() for cell in r.cells]
                    data.append(dict(zip(headers, values)))  # ✅ Map original headers

                return data  # ✅ Returns a list of dictionaries preserving original keys
    
    return None  # ❌ Table not found


def validate_document_revision_history(doc, config):
    """Validates 'Document Revision History' table based on config."""
    mismatches = []
    revision_table = extract_document_revision_history_from_table(doc)

    if not revision_table:
        mismatches.append("Missing 'Document Revision History' table in Document")
        return mismatches

    # ✅ Check if expected columns exist
    expected_columns = {key.replace("DocumentRevisionHistory_", "").lower() for key in config if key.startswith("DocumentRevisionHistory_")}
    print(expected_columns)
    table_columns = set(revision_table[0].keys())  # Extract actual columns from first row
    print(table_columns)

    missing_columns = expected_columns - table_columns
    if missing_columns:
        mismatches.append(f"Missing columns in 'Document Revision History': {', '.join(missing_columns)}")

    # ✅ Validate latest revision date
    revision_dates = [row.get("revision_date", "") for row in revision_table if row.get("revision_date")]
    valid_dates = []

    for date_str in revision_dates:
        try:
            revision_date = datetime.strptime(date_str, "%m/%d/%Y")  # Adjust format if needed
            valid_dates.append(revision_date)
        except ValueError:
            mismatches.append(f"Invalid date format in 'Document Revision History': {date_str}")

    if valid_dates:
        latest_date = max(valid_dates)
        if latest_date < datetime.now():
            mismatches.append(f"Latest revision date {latest_date.strftime('%m/%d/%Y')} is not recent.")

    return mismatches

def validate_table_of_content(doc, config):
    """Check if 'Table of Content' exists in text or tables."""
    toc_found = False
    mismatches = []

    # ✅ Validate 'Document Revision History' separately
    revision_mismatches = validate_document_revision_history(doc, config)
    mismatches.extend(revision_mismatches)

    # Search in all paragraphs
    toc_text = " ".join([p.text.lower() for p in doc.paragraphs])

    if "table of contents" in toc_text:
        toc_found = True
        for key in config:
            if key.startswith("TableOfContent_") and config[key].lower() == "yes":
                expected_value = key.replace("TableOfContent_", "").lower()
                if expected_value not in toc_text:
                    mismatches.append(f"Missing '{expected_value}' in Table of Content")

    # If not found in paragraphs, search in tables
    if not toc_found:
        for table in doc.tables:
            for row in table.rows:
                if any("table of contents" in cell.text.lower() for cell in row.cells):
                    toc_found = True
                    break

    if not toc_found:
        mismatches.append("Table of Content not found")

    return mismatches

def check_embedded_excels(docx_path):
    """Check if the Word document contains multiple embedded Excel files."""
    embedded_excels = []

    try:
        with zipfile.ZipFile(docx_path, "r") as docx_zip:
            for file in docx_zip.namelist():
                if "embeddings" in file.lower() and file.endswith((".xls", ".xlsx", ".xlsm")):
                    embedded_excels.append(file)
    except zipfile.BadZipFile:
        print("❌ The document is not a valid .docx file. Please check the format.")

    return embedded_excels


def validate_sections(extracted_sections, expected_sections):
    """Compare extracted section names with expected sections."""
    extracted_normalized = {s.lower().strip() for s in extracted_sections}
    expected_normalized = {s.lower().strip() for s in expected_sections}

    missing_sections = [s for s in expected_sections if s.lower().strip() not in extracted_normalized]
    extra_sections = [s for s in extracted_sections if s.lower().strip() not in expected_normalized]

    return missing_sections, extra_sections



# def check_embedded_objects(docx_path):
#     """Check if the Word document contains embedded Excel files."""
#     embedded_excels = []

#     try:
#         with zipfile.ZipFile(docx_path, "r") as docx_zip:
#             for file in docx_zip.namelist():
#                 print(file)
#                 if "embeddings" in file.lower() and file.endswith((".xls", ".xlsx", ".xlsm", ".csv")):
#                     embedded_excels.append(os.path.basename(file))
#     except zipfile.BadZipFile:
#         print("❌ The document is not a valid .docx file. Please check the format.")

#     return embedded_excels


import os
import pandas as pd
import zipfile

def extract_embedded_excel(docx_path):
    """Extracts all embedded Excel files from a Word document."""
    embedded_excels = []
    
    try:
        with zipfile.ZipFile(docx_path, "r") as docx_zip:
            for file in docx_zip.namelist():
                if "embeddings" in file.lower() and file.endswith((".xls", ".xlsx", ".xlsm")):
                    output_path = f"{os.path.basename(file)}"
                    
                    with docx_zip.open(file) as src, open(output_path, "wb") as dest:
                        dest.write(src.read())

                    embedded_excels.append(output_path)

    except zipfile.BadZipFile:
        print("❌ The document is not a valid .docx file. Please check the format.")

    return embedded_excels

def validate_excel_content(excel_path, config):
    """Validate that the first sheet contains 'ProjectID' in Row 4 and 'ReleaseID' in Row 6,
       and ensure required sheet names exist.
    """
    try:
        xls = pd.ExcelFile(excel_path)
        sheet_names = [name.lower() for name in xls.sheet_names]  # Normalize for comparison
        required_sheets = {"sheet1", "sheet2", "sheetname"}

        print(f"🔍 Checking {excel_path}...")
        print(f"📄 Found Sheets: {xls.sheet_names}")

        # ✅ Check for required sheets
        missing_sheets = required_sheets - set(sheet_names)
        if missing_sheets:
            reason = f"❌ Missing Sheets: {', '.join(missing_sheets)}"
            print(reason)
            # return False, reason

        # ✅ Read first sheet
        df = pd.read_excel(xls, sheet_name=0, header=None)

        # ✅ Extract expected values from config
        expected_project_id = str(config.get("Page_1_ProjectID", "")).strip().lower()
        expected_release_id = str(config.get("Page_1_ReleaseID", "")).strip().lower()

        # ✅ Ensure we have enough rows before checking
        if df.shape[0] < 6:
            reason = "❌ Excel does not have enough rows for validation."
            print(reason)
            return False, reason

        # ✅ Convert to string (handle NaN cases)
        row_4_value = str(df.iloc[3, 0]).strip().lower()
        row_6_value = str(df.iloc[5, 0]).strip().lower()

        if row_4_value != expected_project_id:
            reason = f"❌ Row 4 Mismatch: Expected '{expected_project_id}', Found '{row_4_value}'"
            print(reason)
            return False, reason

        if row_6_value != expected_release_id:
            reason = f"❌ Row 6 Mismatch: Expected '{expected_release_id}', Found '{row_6_value}'"
            print(reason)
            return False, reason

        print(f"✅ Validation Passed: '{expected_project_id}' in Row 4 & '{expected_release_id}' in Row 6")
        return True, "Validation Passed"

    except Exception as e:
        reason = f"⚠️ Error reading Excel file: {e}"
        print(reason)
        return False, reason


import zipfile
import xml.etree.ElementTree as ET

def get_real_embedded_filename(docx_path, extracted_filename):
    """Retrieve the real name of an embedded Excel file given its extracted filename, with debugging."""
    try:
        with zipfile.ZipFile(docx_path, "r") as docx_zip:
            print("\n🔍 Extracting relationships from document.xml.rels...")
            
            # ✅ Step 1: Extract relationships from document.xml.rels
            rels_path = "word/_rels/document.xml.rels"
            file_rel_map = {}

            if rels_path in docx_zip.namelist():
                with docx_zip.open(rels_path) as rels_file:
                    tree = ET.parse(rels_file)
                    root = tree.getroot()
                    
                    namespace = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
                    for rel in root.findall("rel:Relationship", namespace):
                        target = rel.get("Target")
                        rid = rel.get("Id")
                        if target and target.startswith("embeddings/"):
                            file_name = target.split("/")[-1]  # Extract only filename
                            file_rel_map[file_name] = rid  # Map filename to rId

            print(f"🔹 Extracted file-rel map: {file_rel_map}")

            # ✅ Ensure the extracted filename exists
            if extracted_filename not in file_rel_map:
                print(f"❌ Extracted filename '{extracted_filename}' not found in relationships.")
                return None

            related_rid = file_rel_map[extracted_filename]  # Get associated rId
            print(f"🔹 Found relationship ID: {related_rid}")

            # ✅ Step 2: Extract original names from docProps/app.xml
            app_xml_path = "docProps/app.xml"
            if app_xml_path in docx_zip.namelist():
                print("\n🔍 Extracting TitlesOfParts from docProps/app.xml...")
                with docx_zip.open(app_xml_path) as app_file:
                    tree = ET.parse(app_file)
                    root = tree.getroot()

                    titles_element = root.find(".//{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}TitlesOfParts")

                    if titles_element is not None:
                        title_list = [title.text for title in titles_element]

                        print(f"🔹 Extracted TitlesOfParts: {title_list}")

                        # Convert rId to index
                        rid_index = int(related_rid.replace("rId", "")) - 1

                        if rid_index < len(title_list):
                            real_name = title_list[rid_index]
                            print(f"✅ Real Embedded File Name: {real_name}")
                            return real_name
                        else:
                            print(f"❌ rId index {rid_index} is out of range in TitlesOfParts.")
                    else:
                        print("❌ TitlesOfParts element not found in docProps/app.xml.")

    except Exception as e:
        print(f"❌ Error retrieving real embedded filename: {e}")

    return None  # No match found


def extract_and_validate_embedded_excels(doc_path, config):
    """Extracts all embedded Excel files and validates them, returning detailed failure reasons."""
    embedded_excels = extract_embedded_excel(doc_path)
    validation_results = []

    if not embedded_excels:
        print("❌ No embedded Excel found.")
        return False, "No embedded Excel found"

    all_success = True  # Track if all checks pass

    for excel_file in embedded_excels:
        is_valid, message = validate_excel_content(excel_file, config)
        excel_file = get_real_embedded_filename(doc_path, excel_file)
        validation_results.append(f"📄 {excel_file}: {message}")
        if not is_valid:
            all_success = False

    print("\n✅ Embedded Excel Validation Summary:")
    for result in validation_results:
        print(result)

    if all_success:
        return True, "All embedded Excels validated successfully"
    else:
        return False, "Some embedded Excels failed validation"


def validate_document(doc_path, excel_path):
    """Main function to validate the document using config from Excel."""
    
    # Convert file name to match Excel sheet format
    file_name = os.path.basename(doc_path).replace(".docx", "").replace("-", "_")

    print(f"🔍 Looking for sheet: {file_name} in {excel_path}")

    # Load Excel sheet that matches the modified filename
    config = load_config_from_excel(excel_path, file_name)

    doc = docx.Document(doc_path)

    # Extract content
    text_by_page = extract_text_by_page(doc)
    extracted_sections = extract_section_names(doc)
    tables = extract_table_content(doc)

    # Perform validations
    page_1_mismatches = validate_page_1_content(text_by_page, doc, config)
    toc_mismatches = validate_table_of_content(doc, config)

    
    # missing_sections, extra_sections = validate_sections(extracted_sections, config["Sections"])
    missing_sections, extra_sections = validate_sections_using_toc(doc, config["Sections"])

    # Check for embedded Excel in the document
    # embedded_excels = check_embedded_objects(doc_path)

    # embedded_excels = extract_embedded_excel(doc_path)

    # if embedded_excels:
    #     validation_result = validate_excel_content(embedded_excels, config)
    #     print(validation_result)
    # else:
    #     print("❌ No embedded Excel found.")

    result, reason = extract_and_validate_embedded_excels(doc_path, config)

    if result:
        print("\n✅ Embedded Excel Validation Passed")
    else:
        print(f"\n❌ Embedded Excel validation FAILED: {reason}")


    # Print Report
    print(f"\n=== Validation Report for {file_name} ===")
    
    print("\n✅ Section Validation:")
    print(f"  - Missing Sections: {missing_sections if missing_sections else 'None'}")
    # print(f"  - Unexpected Sections Found: {extra_sections if extra_sections else 'None'}")

    print("\n✅ Page 1 Content Validation:")
    if page_1_mismatches:
        for key, msg in page_1_mismatches.items():
            print(f"  - {key}: {msg}")
    else:
        print("  - All required fields on Page 1 are correct.")

    print("\n✅ Table of Content Validation:")
    if toc_mismatches:
        for issue in toc_mismatches:
            print(f"  - {issue}")
    else:
        print("  - Table of Content contains all required fields.")

    # print("\n✅ Embedded Excel Validation:")
    # if embedded_excels:
    #     print(f"  - Embedded Excel files found: {embedded_excels}")
    # else:
    #     print(f"  - No embedded Excel files found.")

    print("\n🚀 Validation Complete!")


# Example usage
doc_path = "D:\\Desktop 2024\\PycharmProjects\\RESTAPI\\AutomatedDocumentReview\\performance-testing-strategy.docx"
excel_path = "D:\\Desktop 2024\\PycharmProjects\\RESTAPI\\AutomatedDocumentReview\\config.xlsx"  # Excel file with multiple tabs for different Word documents
sheet_name = "Performance_Testing_Strategy"  # Name of the tab to validate

validate_document(doc_path, excel_path)
